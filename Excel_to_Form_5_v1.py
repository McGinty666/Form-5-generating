# -*- coding: utf-8 -*-
"""
Created on Tue Nov 21 22:24:49 2023

@author: robmc
"""

import pandas as pd
from docx import Document
from copy import deepcopy
import os

def read_excel_data(excel_path):
    df = pd.read_excel(excel_path)
    return df

def read_template_document(template_path):
    return Document(template_path)

def add_column_to_word(document, column_data, column_header):
    # Deepcopy the first paragraph of the first cell to preserve formatting
    first_cell_first_paragraph = deepcopy(document.tables[0].cell(0, 0).paragraphs[0])

    # Clear existing text in the table
    for row in document.tables[0].rows:
        for cell in row.cells:
            cell.paragraphs[0].clear()

    # Set the header
    header_cell = document.tables[0].cell(0, 0).paragraphs[0].add_run(column_header)
    header_cell.bold = True

    # Calculate the number of rows and columns in the Word table
    num_rows = len(column_data)
    num_cols = 2

    # Add the column data to the Word table
    for row_index in range(num_rows):
        row_num, col_num = divmod(row_index, num_cols)
        cell = document.tables[0].cell(row_num, col_num)
        data = column_data[row_index]
        if pd.notna(data):  # Check if the data is not NaN
            cell.paragraphs[0].add_run(str(data)).bold = True

    return document

def create_word_documents(template_document, excel_data, output_folder):
    for column_name in excel_data.columns:
        column_data = excel_data[column_name].tolist()
        new_document = deepcopy(template_document)
        new_document = add_column_to_word(new_document, column_data, column_name)

        output_path = os.path.join(output_folder, f"output_{column_name}.docx")
        new_document.save(output_path)
        print(f"Word document created at {output_path}")

if __name__ == "__main__":
    excel_path = "C00015_Form5_wb_v4.xlsx"
    template_path = "E:/Work/ttachments/Form5_empty.docx"
    output_folder = "E:/Work/"

    excel_data = read_excel_data(excel_path)
    template_document = read_template_document(template_path)
    create_word_documents(template_document, excel_data, output_folder)
